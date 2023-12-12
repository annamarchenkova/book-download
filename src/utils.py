from selenium import webdriver  
from selenium.webdriver.common.by import By
# from selenium.webdriver.support.wait import WebDriverWait
# from selenium.webdriver.support import expected_conditions as EC
# from selenium.webdriver.chrome.service import Service
import time
# import pyautogui
import os
from datetime import datetime 
import shutil
import requests
import urllib
import yaml
import re
from PIL import Image
from tqdm import tqdm
import io
from selenium.common.exceptions import NoSuchElementException
from PIL import UnidentifiedImageError
from docx import Document

from project_dirs import PROJECT_DIR, OUTPUT_DIR

def load_config(cnf_dir=PROJECT_DIR, cnf_name="config.yml"):
    """
    load the yaml file
    """
    config_file = open(os.path.join(cnf_dir, cnf_name))
    return yaml.load(config_file, yaml.FullLoader)

def get_proxy_dict(cnf):
    """Get dict in format {'http': proxy, 'https': full_proxy}.

    Args:
        cnf (dict): congiguration parameters
    """

    # Set the proxy server address and port
    proxy = download_proxy(proxy_url=cnf["url_proxy"])
    proxy_server_port = proxy.split("://")[1]

    # Set the proxy username and password
    proxy_username = cnf["proxy_ut"]
    proxy_password = (
        open(os.path.join(PROJECT_DIR, "keys", cnf["proxy_psw_file"]), "r")
        .read()
        .strip("/n")
    )
    # Encode special characters
    proxy_password = urllib.parse.quote(proxy_password)

    # Set the environment variables
    full_proxy = (
        f"http://{proxy_username}:{proxy_password}@{proxy_server_port}"
    )
    os.environ["HTTP_PROXY"] = full_proxy
    os.environ["HTTPS_PROXY"] = full_proxy

    proxy_dict = {"http": full_proxy, "https": full_proxy}

    return proxy_dict


def download_proxy(proxy_url):
    """
    Return the proxy.

    Args:
    url (str): url to downoad the organisational proxy

    Returns:
    (str): proxy value
    """
    text = requests.get(proxy_url, verify=False).text
    proxy = re.findall(r"PROXY (proxy.*)\;\"", text)[0].replace(";", "")

    return "http://" + proxy

def get_webdriver(proxy_address=None):
    """Create webdriver instance

    Args:
        download_dir (str): default download directory

    Returns:
        _type_: webdriver instance
    """
    options = webdriver.ChromeOptions()

    prefs = {
        # "download.default_directory": download_dir,
        "download.prompt_for_download": True,
    }
    options.add_experimental_option("prefs", prefs)
    options.add_argument("--safebrowsing-disable-download-protection")
    options.add_argument("safebrowsing-disable-extension-blacklist")
    if proxy_address:
        options.add_argument(f'--proxy-server={proxy_address}')

    driver = webdriver.Chrome(options=options)
    return driver

def move_files(
    source_path,
    destination_path,
    download_time_threshold,
    extention=".eml",
):
    """Move files with partikcular extention 'extention' that were downloaded
    within the 'download_time_threshold' to 'destination_path'.

    Args:
        source_path (str): path from where move files.
        destination_path (str): where to move files.
        download_time_threshold (datetime.datetime): threshold datetime.
        extention (str, optional): _description_. Defaults to '.eml'.
    """
    for filename in os.listdir(source_path):
        if filename.endswith(extention):
            file_path = os.path.join(source_path, filename)
            # Check if the file was modified in the last hour
            modified_time = datetime.fromtimestamp(
                os.path.getmtime(file_path)
            )
            if modified_time > download_time_threshold:
                # Move the file to the destination folder
                shutil.move(
                    file_path, os.path.join(destination_path, filename)
                )
                # print(f"Moved {filename} to {destination_path}")

def open_page(page_url, driver):
    """Open pec manager webpage

    Args:
        page_url (str): url to open
        driver (webdriver.Chrome): webdriver.Chrome instance

    Returns:
        _type_: _description_
    """
    # apre PEC manager

    driver.get(page_url) 
    driver.maximize_window()
    time.sleep(2)

    return driver


def create_filepath(common_dir, file_name):
    return os.path.join(common_dir, file_name)

def get_book_texts(driver):
    """Get tests from all web pages of the book

    Args:
        driver (webdriver): webdriver instance
    """
    texts = []
    while True:
        page_text = driver.find_element(By.TAG_NAME, 'body').text
        texts.append(page_text)
        try:
            next_page = driver.find_element(
                            By.XPATH, '//*[@id="content-navigation"]/nav/section/div[3]/a/span[2]'
                        )
            next_page.click()
            time.sleep(2) 
        except NoSuchElementException:
            break
    return texts

def save_texts_as_docx_file(texts, common_dir, file_name):
    """Save texts as paragraphs of the docx file.

    Args:
        texts (list of str): texts to save in the final docx
        common_dir (str): common_dir
        file_name (str): file_name without the extention.
    """

    file_path = create_filepath(common_dir=common_dir, file_name=f"{file_name}.docx")

    doc = Document()
    doc.add_heading(file_name, level=1)
    for text in texts:
        heading = text.split("\n")[0]
        paragraphs = text.split("\n")[3:]

        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            sentences = paragraph.split('.')
            if len(sentences) == 1:
                doc.add_heading(paragraph, level=2)
            else:
                doc.add_paragraph(paragraph)

    doc.save(file_path)

def save_book_content_to_docx(driver, file_name, proxies=None, common_dir=OUTPUT_DIR):
    """Get book content and save it to the docx file

    Args:
        driver (webdriver): webdriver instance
        file_name (str): file_name without the extention
    """

    file_path = create_filepath(common_dir=common_dir, file_name=f"{file_name}.docx")

    doc = Document()
    section = doc.sections[0]
    page_width = section.page_width/2
    doc.add_heading(file_name, level=1)

    while True:
        try:
            page_text = driver.find_element(By.TAG_NAME, 'body').text
        except NoSuchElementException:
            page_text = ''

        try:
            # Find all img elements on the page
            img_elements = driver.find_elements(By.TAG_NAME, 'img')

            # Loop through each img element
            for img_element in img_elements:
                img_url = img_element.get_attribute('src')

                try:
                    if proxies:
                        # Download the image
                        response = requests.get(
                            img_url, proxies=proxies, verify=False, timeout=100
                            )
                    else:
                        response = requests.get(img_url, verify=False, timeout=100)
                    image = Image.open(io.BytesIO(response.content))
                    image_path = os.path.join(OUTPUT_DIR, "screenshot.png")
                    image.save(image_path)
                    
                    doc.add_picture(image_path, width=page_width)

                except UnidentifiedImageError:
                    pass

        except NoSuchElementException:
            pass

        if page_text != '':
            heading = page_text.split("\n")[0]
            doc.add_heading(heading, level=1)

            paragraphs = page_text.split("\n")[3:]
            for paragraph in paragraphs:
                sentences = paragraph.split('.')
                if len(sentences) == 1:
                    doc.add_heading(paragraph, level=2)
                else:
                    doc.add_paragraph(paragraph)

        try:
            next_page = driver.find_element(
                            By.XPATH, '//*[@id="content-navigation"]/nav/section/div[3]/a/span[2]'
                        )
            next_page.click()
            time.sleep(2) 
        except NoSuchElementException:
            break

    doc.save(file_path)
    print(f"Saved {file_name}")

def get_title_from_link(link_str):
    """Extract book title from the link

    Args:
        link (str): url
    """
    return link_str.split('/')[5].capitalize().replace('-', ' ')

def open_start_page(proxy, sleep_time=120,
                    link="https://learning.oreilly.com/library/view/introduction-to-machine/9781449369880/preface01.html"):
    """open the first page, logit to the oreilly 

    Args:
        link (str, optional): _description_. Defaults to 
        "https://learning.oreilly.com/library/view/introduction-to-machine/9781449369880/preface01.html".
        link to the first page of the book.
    """
    driver = get_webdriver(proxy_address=proxy)
    driver.get(link)
    time.sleep(sleep_time)
    return driver

def download_from_links(page_urls, proxy, common_dir=OUTPUT_DIR, sleep_time=70):
    """Scrapes the book using the link

    Args:
        page_urls (list of str): List of links for the books
    """
    
    driver = get_webdriver(proxy_address=proxy)
    driver.get("https://learning.oreilly.com/library/view/introduction-to-machine/9781449369880/preface01.html")
    time.sleep(sleep_time)
    
    os.environ["HTTPS_PROXY"] = proxy

    for book_url in tqdm(page_urls):
        # Open a new tab (this opens a blank page by default)
        driver.execute_script("window.open('', '_blank');")
        driver.switch_to.window(driver.window_handles[1])
        driver.get(book_url)
        time.sleep(1)

        book_title = get_title_from_link(book_url)

        # Get book texts and save them as docx
        proxies = {'http': proxy, 'https': proxy}
        save_book_content_to_docx(driver=driver, file_name=book_title, proxies=proxies, common_dir=common_dir)

        # Close the first tab that was opened previously
        driver.switch_to.window(driver.window_handles[0])
        driver.close()

        # Switch back to the new tab
        driver.switch_to.window(driver.window_handles[0])

    driver.quit()
